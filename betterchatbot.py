import sys
sys.path.append(r"C:\Users\bjoco\OneDrive\Dokumen\chatbot\Libarys") # Adjust the path as needed

from adjectives import positive, negative, neutral, hello, bye, Interests
from datetime import datetime
from docx import Document
import glob, random, os, pyttsx3, requests

#debug
start_time = datetime.now()
user_message_count = 0
bot_reply_count = 0
end_after_bye = True

#TTS
Use_TTS = False
selected_voice_index = 0


Name = ""
game_runs = True
DisplayName = "You: "
star_sign = ""
star_sign_name = ""
Bot_Name = "Bot"




# Create session ID and filename at launch
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
session_id = f"User_{timestamp}"
filename = f"chat_archive_{session_id}.docx"

# Initialize the document
doc = Document()
doc.add_heading("Chat Archive", 0)
doc.add_paragraph(f"Session ID: {session_id}")
doc.add_paragraph(f"Session started on: {datetime.now().strftime('%A, %d %B %Y at %I:%M %p')}\n")
doc.save(filename)



def end_chat(): #exits the chat
    responses = [
        f"{random.choice(bye).capitalize()}! buddie",
        f"{random.choice(bye).capitalize()}! it was nice talking to you",
        f"{random.choice(bye).capitalize()}! have a great day",
        f"{random.choice(bye).capitalize()}! hope to chat again soon",
    ]
    print(random.choice(responses))
    if end_after_bye == True:
        exit() 
    else:
        pass

def get_statement(message):
    msg = message.lower()
    if msg == ".name.set":
        return "Name_set"
    elif msg == ".bot.name":
        return "Bot_name"
    elif msg == ".set.tts true":
        global Use_TTS
        Use_TTS = True
        return "TTS_enabled"
    elif msg == ".set.tts false":
        Use_TTS = False
        return "TTS_disabled"
    elif any(word in msg for word in positive):
        return "positive"
    elif msg == ".set.star.sign":
        return "set_star"
    elif msg in ["cmd", "command", "help"]:
        return "help"
    elif any(word in msg for word in negative):
        return "negative"
    elif any(word in msg for word in neutral):
        return "neutral"
    elif any(word in msg for word in Interests):
        return "wow"
    elif any(word in msg.split() for word in hello):
        return "hello"
    elif msg == ".debug.end.after.bye true":
        global end_after_bye
        end_after_bye = True
        return "end_after_bye_true"
    elif msg == ".debug.end.after.bye false":
        end_after_bye = False
        return "end_after_bye_false"
    elif any(word in msg for word in bye):
        end_chat()
        return "bye"
    elif msg == ".debug.bot":
        return "debug"
    elif msg == ".voice.list":
        return "voice_list"
    elif msg.startswith(".voice.set"):
        return "voice_set"
    elif msg == ".clear.archives":
        return "clear_archives"
    elif msg in ["what is your name", "what's your name", "who are you", "who r u", "what are you called", "what are you called?"]:
        return "question_response_name"
    elif msg in ["how r u", "how are you", "how are u", "how r you"]:
        return "question_response_good"
    elif msg in ["what time is it", "what time is it?", "tell me the time", "tell me the time?", "current time", "current time?", "time now", "time now?", "can you tell me the time", "can you tell me the time?", "do you know the time", "do you know the time?", "what's the time", "what's the time?", "what date is it", "what date is it?", "tell me the date", "tell me the date?", "current date", "current date?", "date now", "date now?", "can you tell me the date", "can you tell me the date?", "do you know the date", "do you know the date?", "what's the date", "what's the date?"]:
        return "question_response_time"
    elif msg in ["tell me a joke", "tell me a joke?", "can you tell me a joke", "can you tell me a joke?", "do you know any jokes", "do you know any jokes?"]:
        return "joke"
    elif msg in ["what is my name", "what's my name", "do you know my name", "do you know my name?"]:
        if Name:
            return "user_name_reply"
        else:
            return "Name_set"
    elif msg == ["what do you like", "what do you like?", "what are your interests", "what are your interests?,"  "what do you enjoy", "what do you enjoy?"]:
        return "question_response_interests"
    elif msg == ".weather.update":
        return "update_weather"
    elif "weather" in msg:
        return "tell_weather"
    else:
        return "unknown"

#responses
def generate_response(message):
    statement = get_statement(message)
    
    if statement == "Name_set":
        global Name
        global DisplayName
        Name = input("Enter your name: ")
        DisplayName = Name + star_sign + ": "
        return f"Nice to meet you, {Name}! How can I assist you today?"
    elif statement == "update_weather":
        return update_weather()
    elif statement == "tell_weather":
        update_weather()
        return f"The weather: {current_weather}"
    elif statement == "Bot_name":
        global Bot_Name
        Bot_Name = input("Enter the bot's name: ")
        return f"Got it! I'll go by {Bot_Name} from now on."
    elif statement == "debug":
        current_time = datetime.now()
        return (
        f"Debug Report \n"
        f"Code started running on: {start_time.strftime('%A, %d %B %Y at %I:%M %p')}\n"
        f"User messages sent: {user_message_count}\n"
        f"Bot replies given: {bot_reply_count}\n"
        f"Current date and time: {current_time.strftime('%A, %d %B %Y at %I:%M %p')}"
        )
    elif statement == "TTS_enabled":
        return "Text-to-Speech has been enabled. I will now speak my responses aloud."
    elif statement == "TTS_disabled":
        return "Text-to-Speech has been disabled. I will no longer speak my responses aloud."
    elif statement == "set_star":
        dob_input = input("ENTER DATE OF BIRTH (YYYY-MM-DD): ")
        try:
            dob = datetime.strptime(dob_input, "%Y-%m-%d")
            zodiac_emoji = get_zodiac_sign(dob.month, dob.day)
            DisplayName = f"{Name} {zodiac_emoji}: "
            conversation_log.append(f"🌠 Star Sign Set: {zodiac_emoji} (DOB: {dob_input})")
            return f"Your star sign is {star_sign_name} {zodiac_emoji}, {Name}. May it guide your archive with cosmic resonance."
        except ValueError:
            return "Oops! Please enter your date in YYYY-MM-DD format."
    elif statement == "help":
        return (
            "Available Commands \n"
            ".name.set — Set your name\n"
            ".bot.name — Rename the bot\n"
            ".debug.bot — Show session stats\n"
            ".debug.end.after.bye true/false — Set whether the chat ends after 'bye'\n"
            ".set.tts true/false — Enable or disable Text-to-Speech\n"
            ".weather.update — Update and display current weather info\n"
            ".clear.archives — Clear all chat archives (admin only)\n"
            ".voice.list — List available TTS voices\n"
            ".voice.set [0-2] — Set TTS voice by index from the voice list\n"
            ".set.star.sign — Enter your date of birth to reveal your zodiac sign\n"
            "Type 'bye' to end the chat"
        )
    elif statement == "voice_list":
        list_voices()
        return "Voice list displayed above. Use .voice.set [index] to choose one."
    elif statement == "clear_archives":
        return clear_archives()
    elif statement == "voice_set":
        try:
            index = int(message.split()[-1])
            if index in [0, 1, 2]:
                global selected_voice_index
                selected_voice_index = index
                return f"Voice set to index {index}."
            else:
                return "Only voice indices 0, 1, or 2 are supported right now."
        except ValueError:
            return "Please use the format: .voice.set [0, 1, or 2]"
    elif statement == "end_after_bye_true":
        return "The chat will now end automatically after you say 'bye'."
    elif statement == "end_after_bye_false":
        return "The chat will remain open after you say 'bye'."
    elif statement == "bye":
        pass
    elif statement == "positive":
        responses = [
            f"That's great to hear {Name}! {random.choice(positive).capitalize()} things are always a plus!",
            f"I'm glad you're feeling {random.choice(positive)}! Keep up the good vibes!",
            f"Awesome! {random.choice(positive).capitalize()} experiences make life better!",
            f"Fantastic! It's always {random.choice(positive)} to hear good news!",
        ]
        return random.choice(responses)
    elif statement == "joke":
        jokes = [
            "Why don't scientists trust atoms? Because they make up everything!",
            "Why did the scarecrow win an award? Because he was outstanding in his field!",
            "Why don't skeletons fight each other? They don't have the guts.",
            "What do you call fake spaghetti? An impasta!",
            "Why did the bicycle fall over? Because it was two-tired!",
            "What do you call cheese that isn't yours? Nacho cheese!",
            "Why did the math book look sad? Because it had too many problems.",
            "What do you get when you cross a snowman and a vampire? Frostbite.",
            "Why can't your nose be 12 inches long? Because then it would be a foot!",
            "What do you call a bear with no teeth? A gummy bear!"
        ]
        return random.choice(jokes)
    elif statement == "negative":
        responses = [
            f"I'm sorry to hear that {Name}. Remember, even in {random.choice(negative)} times, there's always a silver lining.",
            f"That's unfortunate. I hope things get better soon {Name}. Stay strong through the {random.choice(negative)} moments.",
            f"That's tough. But remember, every {random.choice(negative)} experience is a chance to learn and grow. Remember that {Name}.",
            f"I understand how you feel {Name}. Sometimes life can be {random.choice(negative)}, but better days are ahead."
        ]
        return random.choice(responses)
    elif statement == "neutral":
        responses = [
            f"Thanks for sharing. It's good to have {random.choice(neutral)} days to balance things out.",
            f"I see. Sometimes {random.choice(neutral)} experiences are just part of life.",
            f"Got it. {random.choice(neutral).capitalize()} moments can be a nice break from the extremes.",
            f"Understood. {random.choice(neutral).capitalize()} days help us appreciate the highs and lows."
        ]
        return random.choice(responses)
    elif statement == "wow":
        responses = [
            f"Yeah i love that its always fun but i also love {random.choice(Interests)}!",
            f"{Name} you like that that so cool have you ever tried {random.choice(Interests)}?",
        ]
        return random.choice(responses)
    
        # simple question responses
    elif statement == "question_response_name":
        responses = [
            f"{random.choice(hello).capitalize()}! How can I assist you today {Name}?",
            f"My name is {Bot_Name}, your friendly chatbot assistant. How can I help you today {Name}?",
            f"I am {Bot_Name}, here to chat and assist you with anything you need {Name}.",
            f"You can call me {Bot_Name}. I'm here to help you with whatever you need {Name}.",
            f"I'm {Bot_Name}, but you can call me anything by typing .bot.name {Name}?"
        ]
        return random.choice(responses)
    elif statement == "user_name_reply":
        responses = [
            f"Your name is {Name}.",
            f"You told me your name is {Name}.",
            f"I remember you said your name is {Name}.",
            f"You mentioned your name is {Name}.",
        ]
        return random.choice(responses)
    elif statement == "question_response_interests":
        responses = [
            f"I enjoy {random.choice(Interests)}! What about you {Name}?",
            f"I'm really into {random.choice(Interests)}. Do you like that too {Name}?",
            f"{random.choice(Interests).capitalize()} is one of my favorite things! How about you {Name}?",
            f"I find {random.choice(Interests)} fascinating! What are your interests {Name}?",
            f"{random.choice(Interests).capitalize()} is something I really enjoy! Do you have any hobbies {Name}?",
        ]
        return random.choice(responses)
    elif statement == "question_response_good":
        responses = [
            f"I am very good thank you for asking {Name}, how are you?",
            f"I am a bot so I don't have feelings but thank you for asking {Name}, how are you?",
            f"I am functioning as expected {Name}, how are you?",
            f"I am doing well thank you {Name}, how are you?",
            f"Good thanks for asking {Name}, how are you?"
        ]
        return random.choice(responses)
    elif statement == "question_response_time":
        current_time = datetime.now().strftime("%A, %d %B %Y at %I:%M %p")
        return f"The current date and time is: {current_time}"      
    

    elif statement == "hello":
        responses = [
            f"{random.choice(hello).capitalize()}! How can I assist you today {Name}?",
            f"{random.choice(hello).capitalize()}! It's great to hear from you {Name}!",
            f"{random.choice(hello).capitalize()}! What would you like to talk about {Name}?",
            f"{random.choice(hello).capitalize()}! I'm here to help with anything you need {Name}.",
            f"Ahh {Name} {random.choice(hello).capitalize()}!"
        ]
        return random.choice(responses) 
    else:
        responses = [
            "I'm not sure how to respond to that, but I'm here to listen!",
            "That's interesting! Can you tell me more?",
            "I see. Feel free to share more about it!",
            "Thanks for sharing! I'm here if you want to talk more."
        ]
        return random.choice(responses)

def get_zodiac_sign(month, day):
    global star_sign
    global star_sign_name
    if (month == 1 and day >= 20) or (month == 2 and day <= 18):
        star_sign = "♒"
        star_sign_name = "Aquarius"
    elif (month == 2 and day >= 19) or (month == 3 and day <= 20):
        star_sign = "♓"
        star_sign_name = "Pisces"
    elif (month == 3 and day >= 21) or (month == 4 and day <= 19):
        star_sign = "♈"
        star_sign_name = "Aries"
    elif (month == 4 and day >= 20) or (month == 5 and day <= 20):
        star_sign = "♉"
        star_sign_name = "Taurus"
    elif (month == 5 and day >= 21) or (month == 6 and day <= 20):
        star_sign = "♊"
        star_sign_name = "Gemini"
    elif (month == 6 and day >= 21) or (month == 7 and day <= 22):
        star_sign = "♋"
        star_sign_name = "Cancer"
    elif (month == 7 and day >= 23) or (month == 8 and day <= 22):
        star_sign = "♌"
        star_sign_name = "Leo"
    elif (month == 8 and day >= 23) or (month == 9 and day <= 22):
        star_sign = "♍"
        star_sign_name = "Virgo"
    elif (month == 9 and day >= 23) or (month == 10 and day <= 22):
        star_sign = "♎"
        star_sign_name = "Libra"
    elif (month == 10 and day >= 23) or (month == 11 and day <= 21):
        star_sign = "♏"
        star_sign_name = "Scorpio"
    elif (month == 11 and day >= 22) or (month == 12 and day <= 21):
        star_sign = "♐"
        star_sign_name = "Sagittarius"
    else:
        star_sign = "♑"
        star_sign_name = "Capricorn"
    return star_sign

conversation_log = []

def clear_archives():
    password = input("Enter password to clear archives: ")
    if password == "8382.clear.archives.admin.2838":
        deleted_files = 0
        for file in glob.glob("*.docx"):
            try:
                os.remove(file)
                deleted_files += 1
                doc = Document()
                doc.add_heading("Admin Chat Archive", 0)
                doc.add_paragraph(f"Session started on: {datetime.now().strftime('%A, %d %B %Y at %I:%M %p')}\n")
                doc.save(filename)
            except Exception as e:
                print(f"Error deleting {file}: {e}")
        return f"Cleared {deleted_files} archive(s) from the folder. fresh archive created."
    else:
        return "incorrect password. Archives not cleared."



def save_conversation_to_doc():
    doc = Document(filename)
    

    for line in conversation_log:
        doc.add_paragraph(line)
    
    current_time = datetime.now()
    doc.add_paragraph("\n Debug Report ")
    doc.add_paragraph(f"Code started running on: {start_time.strftime('%A, %d %B %Y at %I:%M %p')}")
    doc.add_paragraph(f"User messages sent: {user_message_count}")
    doc.add_paragraph(f"Bot replies given: {bot_reply_count}")
    doc.add_paragraph(f"Current date and time: {current_time.strftime('%A, %d %B %Y at %I:%M %p')}")
    
    doc.save(filename)


def list_voices():
    engine = pyttsx3.init()
    voices = engine.getProperty('voices')
    print("Available Voices:")
    for index, voice in enumerate(voices):
        print(f"{index}: {voice.name} — {voice.id}")

def TTS_function(response_text):
    if Use_TTS:
        engine = pyttsx3.init()
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[selected_voice_index].id)
        engine.say(response_text)
        engine.runAndWait()
        engine.stop()





import requests

current_weather = "Weather not set yet."
current_location = None

def update_weather():
    global current_weather, current_location
    api_key = "07c2df9f995e47a1a5a35840251309"


    if not current_location:
        location = input("Enter your city or location: ")
        current_location = location
    else:
        change = input(f"Current location is '{current_location}'. Change it? (y/n): ").strip().lower()
        if change == "y":
            location = input("Enter new city or location: ")
            current_location = location
        else:
            location = current_location

    url = f"http://api.weatherapi.com/v1/current.json?key={api_key}&q={location}&aqi=no"

    try:
        response = requests.get(url)
        data = response.json()

        temp = data["current"]["temp_c"]
        condition = data["current"]["condition"]["text"]
        wind = data["current"]["wind_kph"]

        current_weather = f"{condition}, {temp}°C, wind {wind} km/h"
        return f"✅ Weather updated for {location}: {current_weather}"
    except Exception as e:
        return f"⚠️ Could not update weather: {e}"







def main():
    global user_message_count, bot_reply_count
    print(">>>>>Welcome to the Better Chatbot! Type '.name.set' to set your name and '.bot.name'. Type '.debug.bot' for session stats.<<<<<")
    while game_runs:
        user_input = input(DisplayName)
        user_message_count += 1

        response = generate_response(user_input)
        bot_reply_count += 1

        print(f"{Bot_Name}: {response}")
        if Use_TTS == True:
             TTS_function(response)
        
        conversation_log.append(f"{DisplayName}{user_input}")
        conversation_log.append(f"{Bot_Name}: {response}")
        save_conversation_to_doc()

main()

